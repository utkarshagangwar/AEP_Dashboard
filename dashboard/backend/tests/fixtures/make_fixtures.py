#!/usr/bin/env python3
"""Generate the binary test fixtures for the chunking test suite.

SOW_CHUNKING_PLAN.md §4. Fixtures are GENERATED, not committed as opaque
binaries, so their structure is reviewable in source and adjustable without
a Word/Acrobat round-trip.

Run from backend/:
    python3 tests/fixtures/make_fixtures.py

Requires requirements-dev.txt (python-docx is already a real dependency;
reportlab is dev-only and used here only to build the .pdf fixture).

Output (git-ignored, regenerate on demand):
    tests/fixtures/generated/structured.docx     H1-H3 + 2 interleaved tables
    tests/fixtures/generated/wide_table.docx     one table far over max_chars
    tests/fixtures/generated/numbered.md         clause-numbered markdown
    tests/fixtures/generated/flat.txt            no headings at all
    tests/fixtures/generated/meeting.txt         speaker-turn transcript
    tests/fixtures/generated/pathological.txt    one 60k-char paragraph
    tests/fixtures/generated/multipage.pdf       4 pages, numbered headings
"""
from __future__ import annotations

import os

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated")


def _ensure_out_dir() -> str:
    os.makedirs(OUT_DIR, exist_ok=True)
    return OUT_DIR


def make_structured_docx(path: str) -> None:
    """Headings at three levels with tables INTERLEAVED between sections.

    The interleaving is the entire point: it is what exposes the document
    ordering bug (tables emitted after all paragraphs) and what proves a
    table's heading_path resolves to its real section rather than the last
    heading in the file.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Acme Candidate Portal - Statement of Work", level=0)  # Title
    doc.add_paragraph("This document defines the scope of the engagement.")

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph("The portal lets recruiters manage candidates end to end.")

    doc.add_heading("2. Functional Requirements", level=1)
    doc.add_heading("2.1 Candidate List", level=2)
    doc.add_paragraph("The candidate list is the default landing screen.")

    table = doc.add_table(rows=3, cols=3)
    for col, value in enumerate(("ID", "Control", "Behaviour")):
        table.cell(0, col).text = value
    table.cell(1, 0).text = "REQ-001"
    table.cell(1, 1).text = "Bulk delete button"
    table.cell(1, 2).text = "Deletes all selected candidates after confirmation."
    table.cell(2, 0).text = "REQ-002"
    table.cell(2, 1).text = "Status filter dropdown"
    table.cell(2, 2).text = "Filters by Active, Archived, or All."

    doc.add_heading("2.2 Bulk Actions", level=2)
    doc.add_paragraph("Bulk actions apply to every selected row.")
    doc.add_paragraph("Selection persists across pagination.", style="List Bullet")

    doc.add_heading("2.2.1 Confirmation Modal", level=3)
    doc.add_paragraph("A modal confirms any destructive bulk action.")

    second = doc.add_table(rows=2, cols=2)
    second.cell(0, 0).text = "Field"
    second.cell(0, 1).text = "Required"
    second.cell(1, 0).text = "Reason for deletion"
    second.cell(1, 1).text = "Yes"

    doc.add_heading("3. Sign-off & Acceptance Criteria", level=1)
    doc.add_paragraph("Acceptance requires a full regression pass.")
    doc.save(path)


def make_wide_table_docx(path: str) -> None:
    """A single table well over any sane max_chars, to force row-boundary
    splitting with header repetition (T-C-007)."""
    from docx import Document

    doc = Document()
    doc.add_heading("4. Data Dictionary", level=1)
    rows = 300
    table = doc.add_table(rows=rows + 1, cols=3)
    for col, value in enumerate(("Field", "Type", "Description")):
        table.cell(0, col).text = value
    for r in range(1, rows + 1):
        table.cell(r, 0).text = f"field_{r:03d}"
        table.cell(r, 1).text = "string"
        table.cell(r, 2).text = (
            f"Description for field {r:03d}. " + ("Padding text to grow the row. " * 4)
        )
    doc.save(path)


def make_numbered_md(path: str) -> None:
    sections = [
        "# Acme Portal SOW\n\nIntroductory paragraph for the whole document.\n",
        "## 1. Scope of Work\n\nThe scope covers the candidate management module.\n",
        "### 1.1 Candidate List\n\n"
        "| ID | Control | Notes |\n|---|---|---|\n"
        "| REQ-001 | Bulk delete button | Requires confirmation |\n"
        "| REQ-002 | Status filter | Active / Archived / All |\n",
        "### 1.2 Search\n\n"
        "- Free-text search box\n"
        "- Saved search dropdown\n"
        "  - Rename saved search\n",
        "## 2. Out of Scope\n\nMobile applications are excluded.\n",
        "Setext Heading Style\n====================\n\nContent under a setext heading.\n",
        "### 2.1 Integration Notes\n\n"
        "```python\n# This fence must never be split.\ndef handler():\n    return True\n```\n",
    ]
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(sections))


def make_flat_txt(path: str) -> None:
    """No headings, no structure -- forces the paragraph fallback strategy."""
    paragraph = (
        "The system shall allow the operator to review each record in turn and "
        "apply the appropriate disposition before the batch is closed. "
    )
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n\n".join(paragraph * 6 for _ in range(40)))


def make_meeting_txt(path: str) -> None:
    turns = []
    speakers = ("Priya", "Marcus", "Devi")
    for i in range(60):
        speaker = speakers[i % 3]
        turns.append(
            f"[00:{i // 60:02d}:{i % 60:02d}] {speaker}: "
            "So on the candidate list we definitely need a bulk delete button, "
            "and a status filter next to it. The filter should default to Active. "
            f"That was point number {i}."
        )
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n\n".join(turns))


def make_pathological_txt(path: str) -> None:
    """One unbroken 60k-char paragraph -- the only input that may legitimately
    produce strategy="hard_split" (T-C-014)."""
    sentence = "This single sentence never terminates its paragraph and simply continues. "
    body = sentence * (60_000 // len(sentence) + 1)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(body[:60_000])


def make_multipage_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pages = [
        ("1. Project Overview", ["The portal manages candidates end to end."]),
        ("2. Functional Requirements", [
            "REQ-001 Bulk delete button on the candidate list.",
            "REQ-002 Status filter dropdown with three options.",
        ]),
        ("3. Reporting", ["Weekly export is delivered as CSV."]),
        ("4. Sign-off", ["Acceptance requires a full regression pass."]),
    ]
    pdf = canvas.Canvas(path, pagesize=letter)
    for heading, lines in pages:
        y = 720
        pdf.drawString(72, y, heading)
        for line in lines:
            y -= 20
            pdf.drawString(72, y, line)
        pdf.showPage()
    pdf.save()


BUILDERS = {
    "structured.docx": make_structured_docx,
    "wide_table.docx": make_wide_table_docx,
    "numbered.md": make_numbered_md,
    "flat.txt": make_flat_txt,
    "meeting.txt": make_meeting_txt,
    "pathological.txt": make_pathological_txt,
    "multipage.pdf": make_multipage_pdf,
}


def build_all() -> str:
    out = _ensure_out_dir()
    for name, builder in BUILDERS.items():
        target = os.path.join(out, name)
        builder(target)
        print(f"  wrote {name} ({os.path.getsize(target):,} bytes)")
    return out


if __name__ == "__main__":
    print(f"Generating fixtures into {OUT_DIR}")
    build_all()
    print("Done.")
