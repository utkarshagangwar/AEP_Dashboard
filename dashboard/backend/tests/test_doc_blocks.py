"""Phase 1 — structured extraction (T-E-001..007).

SOW_CHUNKING_PLAN.md §3 Phase 1.

NOTE ON T-E-001: the plan originally specified a byte-identical
characterisation test against the previous extract_existing_sow_text().
That test was retired during implementation because the previous behaviour
was a bug -- .docx tables were emitted after ALL paragraphs, relocating
them out of their section (see doc_blocks module docstring). Byte equality
would have pinned that bug permanently. It is replaced by:

  T-E-001a  content equivalence -- nothing gained, nothing lost
  T-E-001b  ordering IS fixed -- the regression guard against a revert

Together these are strictly stronger than the original: they assert both
that no content changed and that the specific defect is gone.
"""
from __future__ import annotations

import pytest

from app.services import doc_blocks
from app.services.doc_blocks import IngestError, block_text, extract_blocks, render_blocks


def _kinds(blocks):
    return [b["kind"] for b in blocks]


def _headings(blocks):
    return [(b["level"], b["text"]) for b in blocks if b["kind"] == "heading"]


def _all_text(blocks):
    return "\n".join(block_text(b) for b in blocks if block_text(b))


# ── T-E-001a / T-E-001b — the retired characterisation test, replaced ────────

def _legacy_docx_text(storage_path: str) -> str:
    """Verbatim copy of the PREVIOUS sow_import._extract_docx_text().

    Kept here, in the test file only, so the content-equivalence assertion
    below compares against what actually shipped rather than against a
    remembered description of it.
    """
    from docx import Document

    doc = Document(storage_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def test_te001a_docx_render_preserves_all_content(fixture_path):
    """T-E-001a: every non-empty line the old extractor produced still
    appears, and no new content is invented. Only ORDER may differ."""
    path = fixture_path("structured.docx")
    legacy_lines = {ln.strip() for ln in _legacy_docx_text(path).split("\n") if ln.strip()}
    new_lines = {
        ln.strip() for ln in render_blocks(extract_blocks(path, "structured.docx")).split("\n")
        if ln.strip()
    }

    assert not (legacy_lines - new_lines), (
        f"content LOST vs legacy extractor: {legacy_lines - new_lines}"
    )
    assert not (new_lines - legacy_lines), (
        f"content INVENTED vs legacy extractor: {new_lines - legacy_lines}"
    )


def test_te001b_docx_tables_appear_in_document_order(fixture_path):
    """T-E-001b: regression guard for the ordering bug.

    In structured.docx the first table sits under '2.1 Candidate List' and
    '3. Sign-off & Acceptance Criteria' comes later. The legacy extractor
    put the table AFTER sign-off; correct order puts it before.
    """
    path = fixture_path("structured.docx")

    legacy = _legacy_docx_text(path)
    assert legacy.index("REQ-001") > legacy.index("3. Sign-off"), (
        "fixture no longer reproduces the legacy ordering bug -- it must, "
        "or this regression guard proves nothing"
    )

    rendered = render_blocks(extract_blocks(path, "structured.docx"))
    assert rendered.index("REQ-001") < rendered.index("3. Sign-off"), (
        "the .docx table-ordering bug has regressed: table content is again "
        "emitted after later sections"
    )


def test_te001b_table_sits_under_its_own_heading(fixture_path):
    """The stronger form: the nearest preceding heading of the table block
    must be its real section, not whatever heading happens to be last."""
    blocks = extract_blocks(fixture_path("structured.docx"), "structured.docx")
    table_index = next(i for i, b in enumerate(blocks) if b["kind"] == "table")
    preceding = [b for b in blocks[:table_index] if b["kind"] == "heading"]
    assert preceding[-1]["text"] == "2.1 Candidate List"


# ── T-E-002 — .docx heading levels ───────────────────────────────────────────

def test_te002_docx_heading_levels(fixture_path):
    headings = _headings(extract_blocks(fixture_path("structured.docx"), "structured.docx"))
    by_text = {text: level for level, text in headings}

    assert by_text["1. Project Overview"] == 1
    assert by_text["2. Functional Requirements"] == 1
    assert by_text["2.1 Candidate List"] == 2
    assert by_text["2.2 Bulk Actions"] == 2
    assert by_text["2.2.1 Confirmation Modal"] == 3
    # Word's "Title" style maps to level 1.
    assert by_text["Acme Candidate Portal - Statement of Work"] == 1


# ── T-E-003 — .docx tables ───────────────────────────────────────────────────

def test_te003_docx_table_header_separated_from_rows(fixture_path):
    blocks = extract_blocks(fixture_path("structured.docx"), "structured.docx")
    tables = [b for b in blocks if b["kind"] == "table"]

    assert len(tables) == 2
    assert tables[0]["header"] == ["ID", "Control", "Behaviour"]
    assert len(tables[0]["rows"]) == 2
    assert tables[0]["rows"][0][0] == "REQ-001"
    assert tables[0]["rows"][1][1] == "Status filter dropdown"


def test_te003_docx_table_cell_count_preserved(fixture_path):
    blocks = extract_blocks(fixture_path("structured.docx"), "structured.docx")
    table = next(b for b in blocks if b["kind"] == "table")
    assert all(len(row) == len(table["header"]) for row in table["rows"])


def test_te003_first_row_kept_as_data_when_it_has_empty_cells(tmp_path):
    """A table whose first row is NOT a full header must not lose that row to
    the header slot -- the chunker repeats headers on continuation chunks, so
    a misidentified header would be duplicated and its data row lost."""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Only this cell"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "second"
    table.cell(1, 1).text = "row"
    path = tmp_path / "partial_header.docx"
    doc.save(str(path))

    block = next(
        b for b in extract_blocks(str(path), "partial_header.docx") if b["kind"] == "table"
    )
    assert block["header"] == []
    assert len(block["rows"]) == 2


# ── T-E-004 — markdown headings ──────────────────────────────────────────────

def test_te004_atx_and_setext_headings(fixture_path):
    headings = _headings(extract_blocks(fixture_path("numbered.md"), "numbered.md"))
    by_text = {text: level for level, text in headings}

    assert by_text["Acme Portal SOW"] == 1
    assert by_text["1. Scope of Work"] == 2
    assert by_text["1.1 Candidate List"] == 3
    assert by_text["Setext Heading Style"] == 1


def test_te004_md_table_and_code_fence_are_single_blocks(fixture_path):
    blocks = extract_blocks(fixture_path("numbered.md"), "numbered.md")

    tables = [b for b in blocks if b["kind"] == "table"]
    assert len(tables) == 1
    assert tables[0]["header"] == ["ID", "Control", "Notes"]
    assert len(tables[0]["rows"]) == 2

    fences = [b for b in blocks if b["kind"] == "code_fence"]
    assert len(fences) == 1
    assert "def handler():" in fences[0]["text"]
    assert fences[0]["text"].count("```") == 2


def test_te004_md_list_nesting(fixture_path):
    blocks = extract_blocks(fixture_path("numbered.md"), "numbered.md")
    items = [b for b in blocks if b["kind"] == "list_item"]
    texts = {b["text"]: b["level"] for b in items}

    assert texts["Free-text search box"] == 0
    assert texts["Rename saved search"] > texts["Saved search dropdown"]


def test_te004_md_does_not_guess_plaintext_headings(tmp_path):
    """Heuristic heading detection is OFF for .md -- real syntax exists, so
    guessing on top of it only produces false section boundaries."""
    path = tmp_path / "x.md"
    path.write_text("ALL CAPS LINE HERE\n\nsome body text\n", encoding="utf-8")
    blocks = extract_blocks(str(path), "x.md")
    assert _headings(blocks) == []


def test_te004_txt_does_detect_numbered_headings(fixture_path, tmp_path):
    path = tmp_path / "x.txt"
    path.write_text(
        "4.3.2 Bulk Actions\n\nThe bulk action bar appears on selection.\n",
        encoding="utf-8",
    )
    headings = _headings(extract_blocks(str(path), "x.txt"))
    assert headings == [(3, "4.3.2 Bulk Actions")]


# ── T-E-005 — pdf pages ──────────────────────────────────────────────────────

def test_te005_pdf_page_breaks_are_explicit_and_contiguous(fixture_path):
    blocks = extract_blocks(fixture_path("multipage.pdf"), "multipage.pdf")
    pages = [b["page"] for b in blocks if b["kind"] == "page_break"]
    assert pages == [1, 2, 3, 4]


def test_te005_pdf_content_follows_its_page_marker(fixture_path):
    blocks = extract_blocks(fixture_path("multipage.pdf"), "multipage.pdf")

    def page_of(needle: str) -> int:
        current = 0
        for block in blocks:
            if block["kind"] == "page_break":
                current = block["page"]
            elif needle in block_text(block):
                return current
        raise AssertionError(f"{needle!r} not found")

    assert page_of("Project Overview") == 1
    assert page_of("REQ-002") == 2
    assert page_of("Sign-off") == 4


def test_te005_pdf_numbered_headings_detected(fixture_path):
    headings = _headings(extract_blocks(fixture_path("multipage.pdf"), "multipage.pdf"))
    texts = [t for _, t in headings]
    assert "1. Project Overview" in texts
    assert "2. Functional Requirements" in texts


def test_te005_pdf_sentence_lines_are_not_headings(fixture_path):
    """Requirement lines ending in a period must stay paragraphs -- promoting
    them would create a bogus section per requirement."""
    headings = _headings(extract_blocks(fixture_path("multipage.pdf"), "multipage.pdf"))
    assert not any("REQ-00" in text for _, text in headings)


def test_te005_known_limitation_numbered_list_item_reads_as_heading(tmp_path):
    """DOCUMENTED LIMITATION, not desired behaviour.

    In .txt/PDF there is no font metadata, so a short numbered line with no
    terminal punctuation is indistinguishable from a numbered heading. We
    resolve it toward 'heading' because a missed section boundary damages
    every fact beneath it, whereas an extra boundary only splits one list.

    This test exists so the behaviour is a recorded decision rather than a
    surprise. If it starts failing, that is a deliberate change -- update
    the decision, do not silently flip the assertion.
    """
    path = tmp_path / "list.txt"
    path.write_text("1. Enable the feature flag\n\nBody text follows.\n", encoding="utf-8")
    headings = _headings(extract_blocks(str(path), "list.txt"))
    assert headings == [(1, "1. Enable the feature flag")]


def test_te005_deep_numbering_is_never_a_list_item(tmp_path):
    """'4.3.2 X' is unambiguous -- the list regex cannot match it."""
    path = tmp_path / "deep.txt"
    path.write_text("4.3.2 Bulk Actions\n\nBody.\n", encoding="utf-8")
    blocks = extract_blocks(str(path), "deep.txt")
    assert not [b for b in blocks if b["kind"] == "list_item"]
    assert _headings(blocks) == [(3, "4.3.2 Bulk Actions")]


# ── T-E-006 — whitespace handling ────────────────────────────────────────────

def test_te006_blank_paragraphs_dropped(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("real content")
    doc.add_paragraph("   ")
    doc.add_paragraph("")
    doc.add_paragraph("more content")
    path = tmp_path / "blanks.docx"
    doc.save(str(path))

    blocks = extract_blocks(str(path), "blanks.docx")
    assert _kinds(blocks) == ["paragraph", "paragraph"]


def test_te006_bom_is_stripped(tmp_path):
    path = tmp_path / "bom.md"
    path.write_text("﻿# Heading One\n\nbody\n", encoding="utf-8")
    assert _headings(extract_blocks(str(path), "bom.md")) == [(1, "Heading One")]


def test_te006_crlf_does_not_change_block_count(tmp_path):
    body = "# H\n\npara one\n\npara two\n"
    lf = tmp_path / "lf.md"
    crlf = tmp_path / "crlf.md"
    lf.write_text(body, encoding="utf-8")
    crlf.write_bytes(body.replace("\n", "\r\n").encode("utf-8"))

    lf_blocks = extract_blocks(str(lf), "lf.md")
    crlf_blocks = extract_blocks(str(crlf), "crlf.md")
    assert _kinds(lf_blocks) == _kinds(crlf_blocks)
    assert _all_text(lf_blocks) == _all_text(crlf_blocks).replace("\r", "")


# ── T-E-007 — failure modes ──────────────────────────────────────────────────

def test_te007_corrupt_docx_raises_user_safe_error(tmp_path):
    path = tmp_path / "broken.docx"
    path.write_bytes(b"this is definitely not a zip archive")

    with pytest.raises(IngestError) as excinfo:
        extract_blocks(str(path), "broken.docx")
    assert "Could not parse .docx" in str(excinfo.value)


def test_te007_unsupported_extension_raises(tmp_path):
    path = tmp_path / "thing.rtf"
    path.write_text("content", encoding="utf-8")
    with pytest.raises(IngestError) as excinfo:
        extract_blocks(str(path), "thing.rtf")
    assert ".rtf" in str(excinfo.value)


def test_te007_empty_document_raises_never_returns_empty_list(tmp_path):
    path = tmp_path / "empty.md"
    path.write_text("   \n\n  \n", encoding="utf-8")
    with pytest.raises(IngestError):
        extract_blocks(str(path), "empty.md")


def test_te007_image_only_docx_suggests_ocr(tmp_path):
    from docx import Document

    doc = Document()
    path = tmp_path / "images.docx"
    doc.save(str(path))
    with pytest.raises(IngestError) as excinfo:
        extract_blocks(str(path), "images.docx")
    assert "OCR" in str(excinfo.value)


def test_te007_oversized_document_refused(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_blocks, "MAX_SOW_CHARS", 100)
    path = tmp_path / "big.md"
    path.write_text("# H\n\n" + ("x" * 500), encoding="utf-8")
    with pytest.raises(IngestError) as excinfo:
        extract_blocks(str(path), "big.md")
    assert "too large" in str(excinfo.value)


def test_te007_missing_file_raises_ingest_error(tmp_path):
    with pytest.raises(IngestError) as excinfo:
        extract_blocks(str(tmp_path / "nope.md"), "nope.md")
    assert "Could not read document" in str(excinfo.value)
