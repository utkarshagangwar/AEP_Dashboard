"""Structured document extraction — the layer beneath chunking.

SOW_CHUNKING_PLAN.md Phase 1. Turns an uploaded document (.docx / .pdf /
.md / .txt) into an ordered list of typed Blocks that preserve the
structure the author actually encoded — heading levels, table boundaries,
page breaks, list nesting — instead of flattening everything to a string.

WHY THIS EXISTS
---------------
app.services.doc_chunking needs to split large documents at *structural*
boundaries (never mid-table, never mid-clause) and to label each chunk with
the section it came from. Neither is possible against a flat string: once
"## 4.3 Candidate Management" is just another line of text, the only thing
a splitter can do is count characters, which is exactly the behaviour this
plan replaces.

DEVIATION FROM THE PLAN'S FILE MANIFEST
--------------------------------------
SOW_CHUNKING_PLAN.md §4 put block extraction inside sow_import.py and
design_ingest.py separately. It lives here instead, as one module, because
both of those modules need the *same* block model and the same .pdf/.md
extractors -- splitting it would have duplicated the parsers or created an
import cycle between the two. sow_import.extract_existing_sow_blocks() and
design_ingest.extract_text() are now thin delegates to this module, so
their public signatures are unchanged for every existing caller.

TWO BUGS THIS FIXES (found during Phase 1, see plan §0)
-------------------------------------------------------
1. .docx DOCUMENT ORDER. The previous sow_import._extract_docx_text()
   iterated `doc.paragraphs` to exhaustion and *then* `doc.tables`, which
   relocated every table to the end of the document. A requirements table
   under "2. Requirements" was emitted after "5. Sign-off", severing it
   from its heading. On a document over one chunk's worth of text, every
   table landed together in the final chunk. Fixed here by walking
   `document.element.body` in true XML order (_iter_docx_block_items).

   This deliberately changes extract_existing_sow_text()'s output. The old
   byte-for-byte behaviour is NOT preserved -- it was wrong. The
   characterisation test asserts content equivalence (nothing gained,
   nothing lost, see tests/test_doc_blocks.py::test_docx_render_content_set)
   rather than byte equality.

2. .pdf PAGE BOUNDARIES. design_ingest.extract_text() joined pages with
   "\\n\\n", making a page break indistinguishable from a paragraph break,
   so no chunk could ever carry a "p.12" locator. Pages are now explicit
   page_break blocks.

RELIABILITY RULE (unchanged from design_ingest/sow_ledger)
----------------------------------------------------------
Extraction failures raise IngestError with a user-safe message. An empty
result is never returned as success.
"""
from __future__ import annotations

import os
import re
from typing import Any, Iterator, Literal, TypedDict

from app.core.logging import get_logger

logger = get_logger(__name__)

# Sanity ceiling only -- guards against pathological input, not normal size.
# Kept identical to the value design_ingest and sow_import already enforce so
# the refusal threshold a user experiences does not shift.
MAX_SOW_CHARS = 2_000_000


class IngestError(RuntimeError):
    """Raised when a document cannot be ingested; message is user-safe.

    Defined here rather than in design_ingest so this module has no import
    dependency on it (design_ingest imports *this* module, not the reverse).
    design_ingest re-exports this exact class object, so every existing
    `except design_ingest.IngestError` and `from app.services.design_ingest
    import IngestError` keeps working unchanged.
    """


# ── Block model ──────────────────────────────────────────────────────────────
#
# A Block is a plain dict (not a dataclass) so it round-trips through JSONB
# without a converter -- sow_parts.heading_path and any future debugging
# payload store these directly.

BlockKind = Literal[
    "heading", "paragraph", "list_item", "table", "page_break", "code_fence"
]


class Block(TypedDict, total=False):
    kind: BlockKind
    text: str            # heading / paragraph / list_item / code_fence
    level: int           # heading: 1-9. list_item: nesting depth, 0-based.
    header: list[str]    # table: header row cells, [] if the table has none
    rows: list[list[str]]  # table: body rows (header excluded)
    page: int            # page_break: 1-based number of the page STARTING here


def _heading(text: str, level: int) -> Block:
    return {"kind": "heading", "level": max(1, min(9, level)), "text": text}


def _paragraph(text: str) -> Block:
    return {"kind": "paragraph", "text": text}


# ── .docx ────────────────────────────────────────────────────────────────────

# python-docx exposes paragraph style names as "Heading 1".."Heading 9".
_DOCX_HEADING_RE = re.compile(r"^heading\s*(\d)$", re.IGNORECASE)


def _iter_docx_block_items(document: Any) -> Iterator[tuple[str, Any]]:
    """Yield ("paragraph", Paragraph) / ("table", Table) in TRUE document
    order by walking the body XML directly.

    python-docx's `document.paragraphs` and `document.tables` are two
    separate flat views with no interleaving information between them --
    reading both in sequence is what produced the ordering bug this module
    exists to fix. The body element is the only place the real order lives.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, document)
        # Anything else (sectPr, bookmarks, etc.) carries no extractable
        # text and is skipped -- same as the previous implementation.


def _docx_heading_level(paragraph: Any) -> int | None:
    """Heading level from the paragraph's style, or None if it is body text.

    Falls back to a shape heuristic for documents built on custom styles
    (very common in client-supplied SOWs, where "SOW Heading 2" is not a
    name python-docx recognises). The fallback is logged at DEBUG so its
    real-world frequency is measurable rather than guessed at -- if it fires
    constantly on real documents, that is the signal to parse the style
    hierarchy properly (plan §6, LOW risk row).
    """
    style_name = ""
    try:
        style_name = (paragraph.style.name or "").strip()
    except (AttributeError, KeyError):
        # A style referencing a definition missing from the .docx -- rare but
        # real in files produced by non-Word editors. Not fatal: fall through
        # to the heuristic rather than failing the whole import.
        pass

    if style_name.lower() == "title":
        return 1
    match = _DOCX_HEADING_RE.match(style_name)
    if match:
        return int(match.group(1))

    # Heuristic fallback: a short, standalone, fully-bold line with no
    # sentence-ending punctuation is a heading in practice.
    text = paragraph.text.strip()
    if not text or len(text) > 120 or text.endswith((".", ",", ";", ":")):
        return None
    runs = [r for r in paragraph.runs if r.text.strip()]
    if runs and all(r.bold for r in runs):
        logger.debug("doc_blocks: custom-style heading detected by heuristic: %r", text[:60])
        # Depth is unknowable from formatting alone; 2 keeps it below any
        # real Title/H1 rather than competing with it.
        return 2
    return None


def _docx_list_level(paragraph: Any) -> int | None:
    """List nesting depth for "List Bullet 2" / "List Number" style names."""
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except (AttributeError, KeyError):
        return None
    if not style_name.startswith("list"):
        return None
    trailing = re.search(r"(\d)\s*$", style_name)
    return int(trailing.group(1)) - 1 if trailing else 0


def _docx_table_block(table: Any) -> Block | None:
    """A table as ONE block. The first row is treated as the header only when
    every one of its cells is non-empty -- a table whose first row already
    holds data would otherwise silently lose that row's content to a header
    slot the chunker then repeats on every continuation chunk.
    """
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return None

    header: list[str] = []
    if len(rows) > 1 and all(c for c in rows[0]):
        header = rows[0]
        rows = rows[1:]
    return {"kind": "table", "header": header, "rows": rows}


def _extract_docx_blocks(storage_path: str) -> list[Block]:
    try:
        from docx import Document

        document = Document(storage_path)
    except Exception as exc:  # noqa: BLE001 — python-docx raises several types
        raise IngestError(f"Could not parse .docx file: {exc}") from exc

    blocks: list[Block] = []
    for kind, item in _iter_docx_block_items(document):
        if kind == "table":
            table_block = _docx_table_block(item)
            if table_block:
                blocks.append(table_block)
            continue

        text = item.text.strip()
        if not text:
            continue

        level = _docx_heading_level(item)
        if level is not None:
            blocks.append(_heading(text, level))
            continue

        list_level = _docx_list_level(item)
        if list_level is not None:
            blocks.append({"kind": "list_item", "level": list_level, "text": text})
            continue

        blocks.append(_paragraph(text))

    return blocks


# ── Markdown / plain text ────────────────────────────────────────────────────

_ATX_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")
_SETEXT_UNDERLINE_RE = re.compile(r"^(=+|-{2,})\s*$")
_FENCE_RE = re.compile(r"^(```|~~~)")
_LIST_ITEM_RE = re.compile(r"^(\s*)(?:[-*+]|\d+[.)])\s+(.*)$")
_PIPE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_PIPE_DIVIDER_RE = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
# A numbered clause heading in a plain-text SOW: "4.3.2 Bulk Actions".
# Requires real following text so a bare list item ("1. do the thing") that
# ends in a period is not mistaken for a section heading.
_NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+(\S.{0,98})$")


def _split_pipe_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _looks_like_plaintext_heading(line: str) -> int | None:
    """Heading level for a plain-text/PDF line, or None.

    Deliberately conservative: a false positive here creates a bogus section
    boundary and mislabels every fact under it, which is worse than missing
    a heading and falling back to paragraph splitting.
    """
    text = line.strip()
    if not text or len(text) > 100 or text.endswith((".", ",", ";")):
        return None

    match = _NUMBERED_HEADING_RE.match(text)
    if match:
        # "4.3.2 Bulk Actions" -> depth 3. Capped at 6 to stay within the
        # heading level range the rest of the pipeline assumes.
        return min(6, match.group(1).count(".") + 1)

    # ALL CAPS section titles ("SCOPE OF WORK") -- require at least two words
    # so an acronym on its own line ("API") is not promoted to a heading.
    letters = [c for c in text if c.isalpha()]
    if letters and all(c.isupper() for c in letters) and len(text.split()) >= 2:
        return 1
    return None


def _extract_markdown_blocks(text: str, *, detect_plaintext_headings: bool) -> list[Block]:
    """Parse markdown-ish text into blocks.

    detect_plaintext_headings enables the numbered/ALL-CAPS heuristic used
    for .txt and PDF-derived text. It is OFF for .md, where real syntax
    exists and guessing on top of it only adds false positives.
    """
    lines = text.split("\n")
    blocks: list[Block] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        joined = "\n".join(paragraph_lines).strip()
        if joined:
            blocks.append(_paragraph(joined))
        paragraph_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block -- captured whole; never split by the chunker.
        fence = _FENCE_RE.match(stripped)
        if fence:
            flush_paragraph()
            marker = fence.group(1)
            fence_lines = [line]
            i += 1
            while i < len(lines):
                fence_lines.append(lines[i])
                if lines[i].strip().startswith(marker):
                    i += 1
                    break
                i += 1
            blocks.append({"kind": "code_fence", "text": "\n".join(fence_lines)})
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        atx = _ATX_RE.match(stripped)
        if atx:
            flush_paragraph()
            blocks.append(_heading(atx.group(2), len(atx.group(1))))
            i += 1
            continue

        # Setext: the underline belongs to the line ABOVE it, which is why
        # this is checked before the current line is buffered as a paragraph.
        if (
            i + 1 < len(lines)
            and _SETEXT_UNDERLINE_RE.match(lines[i + 1].strip())
            and not _LIST_ITEM_RE.match(line)
        ):
            flush_paragraph()
            blocks.append(_heading(stripped, 1 if lines[i + 1].strip()[0] == "=" else 2))
            i += 2
            continue

        # Pipe table -- consumed as one block so the chunker cannot split it.
        if _PIPE_ROW_RE.match(line):
            flush_paragraph()
            table_lines = []
            while i < len(lines) and _PIPE_ROW_RE.match(lines[i]):
                table_lines.append(lines[i])
                i += 1
            rows = [
                _split_pipe_row(raw)
                for raw in table_lines
                if not _PIPE_DIVIDER_RE.match(raw)
            ]
            if rows:
                has_divider = any(_PIPE_DIVIDER_RE.match(raw) for raw in table_lines)
                header = rows[0] if (has_divider and len(rows) > 1) else []
                blocks.append({
                    "kind": "table",
                    "header": header,
                    "rows": rows[1:] if header else rows,
                })
            continue

        # Plaintext heading detection MUST run before list detection.
        # "1. Project Overview" matches _LIST_ITEM_RE ("1." + space), so
        # checking lists first classified every top-level numbered heading in
        # a .txt/PDF SOW as a list item and produced zero sections.
        #
        # The two are genuinely ambiguous without font metadata, which
        # pypdf does not provide. _looks_like_plaintext_heading discriminates
        # on length (<=100) and absence of terminal sentence punctuation,
        # because a numbered LIST item is normally a sentence and a numbered
        # HEADING normally is not. Known limitation, asserted explicitly in
        # tests/test_doc_blocks.py: a short, period-less numbered list item
        # in a .txt or PDF is read as a heading. Deeper numbering ("4.3.2 X")
        # is unambiguous -- _LIST_ITEM_RE never matches it.
        if detect_plaintext_headings and not paragraph_lines:
            level = _looks_like_plaintext_heading(stripped)
            if level is not None:
                flush_paragraph()
                blocks.append(_heading(stripped, level))
                i += 1
                continue

        list_match = _LIST_ITEM_RE.match(line)
        if list_match:
            flush_paragraph()
            indent = len(list_match.group(1).expandtabs(4))
            blocks.append({
                "kind": "list_item",
                "level": indent // 2,
                "text": list_match.group(2).strip(),
            })
            i += 1
            continue

        paragraph_lines.append(line)
        i += 1

    flush_paragraph()
    return blocks


# ── .pdf ─────────────────────────────────────────────────────────────────────

def _extract_pdf_blocks(storage_path: str) -> list[Block]:
    """Page-aware PDF extraction.

    Each page contributes an explicit page_break block before its content,
    which is what makes a "p.12" locator possible. The previous
    "\\n\\n".join(pages) made page boundaries indistinguishable from
    paragraph boundaries and permanently lost that information.

    pypdf's extract_text() gives no font or size metadata, so headings
    inside a page can only be detected heuristically -- hence
    detect_plaintext_headings=True here.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(storage_path)
        if reader.is_encrypted:
            raise IngestError("PDF is password-protected; upload an unlocked copy.")
        pages = [(page.extract_text() or "") for page in reader.pages]
    except IngestError:
        raise
    except Exception as exc:  # noqa: BLE001 — pypdf raises many exception types
        raise IngestError(f"Could not parse PDF: {exc}") from exc

    blocks: list[Block] = []
    for page_number, page_text in enumerate(pages, start=1):
        blocks.append({"kind": "page_break", "page": page_number})
        if page_text.strip():
            blocks.extend(
                _extract_markdown_blocks(page_text, detect_plaintext_headings=True)
            )
    return blocks


# ── Public entry points ──────────────────────────────────────────────────────

DOCX_EXTENSIONS = (".docx",)
TEXT_EXTENSIONS = (".txt", ".md")
PDF_EXTENSIONS = (".pdf",)
SUPPORTED_EXTENSIONS = DOCX_EXTENSIONS + TEXT_EXTENSIONS + PDF_EXTENSIONS


def extract_blocks(storage_path: str, file_name: str) -> list[Block]:
    """Structured blocks for a document, in true document order.

    Raises IngestError for unsupported formats, unreadable files, and
    documents that yield no text -- never returns [] as a success.
    """
    ext = os.path.splitext(file_name.lower())[1]

    if ext in DOCX_EXTENSIONS:
        blocks = _extract_docx_blocks(storage_path)
        empty_hint = (
            "No text could be extracted from this .docx file. If it's mostly "
            "images/scans, it needs OCR first."
        )
    elif ext in PDF_EXTENSIONS:
        blocks = _extract_pdf_blocks(storage_path)
        empty_hint = (
            "No text could be extracted. If this is a scanned PDF, it needs OCR first."
        )
    elif ext in TEXT_EXTENSIONS:
        try:
            with open(storage_path, "r", encoding="utf-8", errors="replace") as fh:
                raw = fh.read()
        except OSError as exc:
            raise IngestError(f"Could not read document: {exc}") from exc
        # BOM would otherwise become part of the first heading's text.
        raw = raw.lstrip("﻿")
        blocks = _extract_markdown_blocks(
            raw, detect_plaintext_headings=(ext == ".txt")
        )
        empty_hint = "No text could be extracted; the file appears to be empty."
    else:
        raise IngestError(
            f"Unsupported document format '{ext}'. Use .docx, .pdf, .txt, or .md."
        )

    if not any(b["kind"] != "page_break" for b in blocks):
        raise IngestError(empty_hint)

    total = sum(len(t) for t in (block_text(b) for b in blocks))
    if total > MAX_SOW_CHARS:
        raise IngestError(
            f"Document is too large to ingest ({total:,} chars). "
            "Split it into smaller files and upload separately."
        )

    logger.info(
        "doc_blocks: extracted %d block(s) (%d chars) from %s", len(blocks), total, file_name
    )
    return blocks


def block_text(block: Block) -> str:
    """The rendered text of a single block.

    One function so the renderer, the chunker's size accounting, and the
    lossless-invariant test all agree on what a block's length is. If these
    ever diverged, chunks could silently exceed max_chars.
    """
    kind = block["kind"]
    if kind == "table":
        lines = []
        if block.get("header"):
            lines.append(" | ".join(block["header"]))
        lines.extend(" | ".join(row) for row in block.get("rows", []))
        return "\n".join(lines)
    if kind == "page_break":
        return ""
    return block.get("text", "")


def render_blocks(blocks: list[Block]) -> str:
    """Blocks back to plain text.

    Output matches the previous flat extractors line-for-line (paragraph
    text as-is, table rows joined with " | ") with ONE intentional
    difference: .docx tables now appear in their real position instead of
    being appended after all paragraphs. See this module's docstring.
    """
    parts = [block_text(b) for b in blocks]
    return "\n".join(p for p in parts if p)
