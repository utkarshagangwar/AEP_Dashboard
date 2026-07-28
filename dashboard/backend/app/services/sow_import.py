"""Import SOW (SOW tab) — text extraction for an uploaded pre-existing SOW.

Feeds app.services.sow_ledger's extract_ledger_from_sow_document* (which in
turn feeds app.workers.tasks.sow_ledger::extract_existing_sow_ledger_task):
a user attaches an already-written SOW/requirements document as a source on
a sow_documents row, this module turns it into plain text, and the ledger
extraction described in SOW_FEATURE_PLAN.md §2 Pass 1 takes it from there —
same downstream contract every other source type (transcript/recording/
design) already uses.

Deliberately a separate module from app.services.design_ingest, not an
extension of it: design_ingest.extract_text() is reused UNMODIFIED here for
.txt/.md/.pdf (identical extraction need, zero reason to fork it), but this
module adds .docx support, which design_ingest has never needed (the
SOW-Checkpoints pipeline it serves has only ever accepted txt/md/pdf) —
adding it there would be an unrelated behavior change to an existing,
already-shipped endpoint. Keeping it here means that pipeline's behavior is
untouched.

Reliability rule, same as design_ingest/sow_ledger: extraction failures
raise IngestError with a user-safe message — never a silent empty result.
"""
from __future__ import annotations

import os

from app.core.logging import get_logger
from app.services.design_ingest import IngestError, extract_text

logger = get_logger(__name__)

# Same ceiling design_ingest.extract_text applies to txt/md/pdf -- kept in
# sync deliberately, not imported (that constant is module-private there),
# so a very large parsed .docx is refused rather than silently truncated,
# consistent with how oversized txt/md/pdf are already handled.
_MAX_SOW_CHARS = 2_000_000

_DOCX_EXTENSIONS = (".docx",)
_DELEGATED_EXTENSIONS = (".txt", ".md", ".pdf")
SUPPORTED_EXTENSIONS = _DOCX_EXTENSIONS + _DELEGATED_EXTENSIONS


def _extract_docx_text(storage_path: str) -> str:
    """Plain text from a .docx: every paragraph, plus every table cell (SOWs
    frequently put functional requirements in tables — dropping them would
    silently lose exactly the kind of exhaustive detail this feature exists
    to preserve, per SOW_FEATURE_PLAN.md's "if anything goes missing... it
    will impact the business" constraint)."""
    try:
        from docx import Document

        doc = Document(storage_path)
    except Exception as exc:  # noqa: BLE001 — python-docx raises several exception types
        raise IngestError(f"Could not parse .docx file: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_existing_sow_text(storage_path: str, file_name: str) -> str:
    """Extract plain text from an uploaded existing-SOW file. Supports
    .docx, .pdf, .txt, .md — the same format set the "Import SOW" upload
    endpoint (app/api/v1/sow.py) validates against."""
    ext = os.path.splitext(file_name.lower())[1]

    if ext in _DOCX_EXTENSIONS:
        text = _extract_docx_text(storage_path).strip()
        if not text:
            raise IngestError(
                "No text could be extracted from this .docx file. If it's mostly "
                "images/scans, it needs OCR first."
            )
        if len(text) > _MAX_SOW_CHARS:
            raise IngestError(
                f"Document is too large to ingest ({len(text):,} chars). "
                "Split it into smaller files and upload separately."
            )
        return text

    if ext in _DELEGATED_EXTENSIONS:
        # Unmodified reuse of the existing SOW-Checkpoints extractor's text
        # extraction -- identical need, zero duplicated logic/behavior.
        return extract_text(storage_path, file_name)

    raise IngestError(
        f"Unsupported SOW format '{ext}'. Use .docx, .pdf, .txt, or .md."
    )
