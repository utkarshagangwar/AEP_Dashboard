"""SOW export rendering (Phase 6 — "Pass 4 render-to-format", per
SOW_FEATURE_PLAN.md §4/§11.7).

Three formats, all rendered ON DEMAND from the same structured
`content_blocks` source of truth every other pass in this feature reads
from — never persisted (plan §11.7): a stored export file would go stale
the instant a section is edited (Phase 5) or patched (Phase 7), which is
exactly the kind of silent-trust problem this whole feature exists to
prevent. Every call to these functions re-renders fresh from whatever is
in the database right now.

  render_document_markdown — trivial, reuses sow_drafting.render_blocks_
  markdown per section. Source-of-truth-adjacent: this is the same
  rendering the in-app read-only view already shows per section.

  render_document_html / render_document_docx — NOT markdown pasted into
  a document. Each block type gets a real native element (a real Word
  table via python-docx, a real <table> for weasyprint to lay out) so the
  exported file looks like an actual document, not a markdown dump with
  visible '#'/'|' characters. Both walk the same typed block list, so a
  bug in one converter's block-type coverage would show up as a visibly
  wrong export, not a silent format mismatch between them.

  render_document_pdf — renders render_document_html's output via
  weasyprint. Needs system Pango/Cairo/GDK-pixbuf libraries in the actual
  deployment container (docker/Dockerfile.backend) — importing/running
  weasyprint in a bare Python environment without them fails at PDF-write
  time, not at import time, so this is deliberately exercised end-to-end
  (not just imported) during this feature's own verification pass.
"""
from __future__ import annotations

import html as html_stdlib
import io
from datetime import datetime, timezone

from app.core.logging import get_logger

logger = get_logger(__name__)

_ELEMENT_TYPE_LABELS = {
    "three_dot_menu": "three-dot menu",
}


def _element_label(element_type: str | None) -> str:
    if not element_type:
        return "control"
    return _ELEMENT_TYPE_LABELS.get(element_type, element_type.replace("_", " "))


# ── Markdown (trivial — reuses the per-section renderer everything else uses) ──

def render_document_markdown(document_title: str, sections: list) -> str:
    """sections: SowSection ORM rows, already ordered by order_index."""
    from app.services.sow_drafting import render_blocks_markdown

    parts = [f"# {document_title}", ""]
    for section in sections:
        parts.append(render_blocks_markdown(section.content_blocks or []))
        parts.append("")
    return "\n".join(parts).strip() + "\n"


# ── HTML (native elements, feeds the PDF renderer) ──────────────────────────

_PDF_CSS = """
@page { size: A4; margin: 2.2cm; }
body { font-family: "Liberation Sans", Arial, sans-serif; color: #1F2937; font-size: 10.5pt; line-height: 1.5; }
h1 { font-size: 20pt; margin: 0 0 6pt; color: #111827; }
h2 { font-size: 14pt; margin: 20pt 0 8pt; color: #111827; border-bottom: 1pt solid #E5E7EB; padding-bottom: 4pt; }
h3 { font-size: 12pt; margin: 14pt 0 6pt; color: #111827; }
h4 { font-size: 11pt; margin: 10pt 0 4pt; color: #111827; }
p { margin: 0 0 8pt; }
ul { margin: 0 0 8pt; padding-left: 18pt; }
li { margin-bottom: 3pt; }
table { border-collapse: collapse; width: 100%; margin: 0 0 10pt; }
th, td { border: 1pt solid #D1D5DB; padding: 5pt 7pt; text-align: left; font-size: 9.5pt; vertical-align: top; }
th { background: #F3F4F6; font-weight: 600; }
.control { margin: 0 0 6pt; padding: 6pt 8pt; background: #F9FAFB; border-left: 3pt solid #2563EB; }
.control .label { font-weight: 600; }
.control .kind { color: #6B7280; font-size: 9pt; text-transform: uppercase; letter-spacing: 0.03em; }
.callout { margin: 0 0 8pt; padding: 6pt 10pt; border-radius: 3pt; font-size: 9.5pt; }
.callout.info { background: #EFF6FF; border: 1pt solid #BFDBFE; }
.callout.warning { background: #FEF3C7; border: 1pt solid #FDE68A; }
.meta { color: #9CA3AF; font-size: 8.5pt; margin: 0 0 16pt; }
"""


def _block_to_html(block: dict) -> str:
    e = html_stdlib.escape
    btype = block.get("type")

    if btype == "heading":
        level = max(1, min(int(block.get("level") or 2), 4))
        return f"<h{level}>{e(block.get('text', ''))}</h{level}>"

    if btype == "paragraph":
        return f"<p>{e(block.get('text', ''))}</p>"

    if btype == "control_spec":
        label = e(block.get("label", ""))
        kind = e(_element_label(block.get("element_type")))
        behavior = e(block.get("behavior") or "")
        behavior_html = f" — {behavior}" if behavior else ""
        return (
            f'<div class="control"><span class="label">{label}</span> '
            f'<span class="kind">({kind})</span>{behavior_html}</div>'
        )

    if btype == "bullet_list":
        items = "".join(f"<li>{e(item)}</li>" for item in block.get("items") or [])
        return f"<ul>{items}</ul>"

    if btype == "table":
        headers = block.get("headers") or []
        rows = block.get("rows") or []
        head_html = "".join(f"<th>{e(h)}</th>" for h in headers)
        body_html = "".join(
            "<tr>" + "".join(f"<td>{e(c)}</td>" for c in row) + "</tr>" for row in rows
        )
        return f"<table><thead><tr>{head_html}</tr></thead><tbody>{body_html}</tbody></table>"

    if btype == "callout":
        tone = block.get("tone") if block.get("tone") in ("info", "warning") else "info"
        icon = "⚠" if tone == "warning" else "ℹ"
        return f'<div class="callout {tone}">{icon} {e(block.get("text", ""))}</div>'

    return ""


def render_document_html(document_title: str, sections: list) -> str:
    body_parts = [f"<h1>{html_stdlib.escape(document_title)}</h1>"]
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    body_parts.append(f'<p class="meta">Generated {generated}</p>')
    for section in sections:
        for block in section.content_blocks or []:
            rendered = _block_to_html(block)
            if rendered:
                body_parts.append(rendered)
    body = "\n".join(body_parts)
    return (
        "<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
        f"<title>{html_stdlib.escape(document_title)}</title>"
        f"<style>{_PDF_CSS}</style></head><body>{body}</body></html>"
    )


def render_document_pdf(document_title: str, sections: list) -> bytes:
    """Raises whatever weasyprint raises on failure (e.g. missing system
    libs in a container that hasn't picked up the Dockerfile change yet) —
    the API layer is responsible for turning that into a clean 500 with a
    server-side log, not this function papering over it."""
    import weasyprint

    html_doc = render_document_html(document_title, sections)
    return weasyprint.HTML(string=html_doc).write_pdf()


# ── DOCX (native python-docx objects) ───────────────────────────────────────

def _add_block_to_docx(doc, block: dict) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    btype = block.get("type")

    if btype == "heading":
        level = max(1, min(int(block.get("level") or 2), 4))
        doc.add_heading(block.get("text", ""), level=level)
        return

    if btype == "paragraph":
        doc.add_paragraph(block.get("text", ""))
        return

    if btype == "control_spec":
        p = doc.add_paragraph()
        run = p.add_run(block.get("label", ""))
        run.bold = True
        kind_run = p.add_run(f"  ({_element_label(block.get('element_type'))})")
        kind_run.italic = True
        kind_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        behavior = block.get("behavior") or ""
        if behavior:
            p.add_run(f" — {behavior}")
        return

    if btype == "bullet_list":
        for item in block.get("items") or []:
            doc.add_paragraph(item, style="List Bullet")
        return

    if btype == "table":
        headers = block.get("headers") or []
        rows = block.get("rows") or []
        if not headers:
            return
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        for c, header_text in enumerate(headers):
            cell = table.rows[0].cells[c]
            cell.text = str(header_text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        for r, row in enumerate(rows, start=1):
            for c, cell_text in enumerate(row):
                if c < len(headers):
                    table.rows[r].cells[c].text = str(cell_text)
        return

    if btype == "callout":
        tone = block.get("tone") if block.get("tone") in ("info", "warning") else "info"
        icon = "⚠" if tone == "warning" else "ℹ"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        run = p.add_run(f"{icon} {block.get('text', '')}")
        run.italic = True
        return


def render_document_docx(document_title: str, sections: list) -> bytes:
    import docx

    doc = docx.Document()
    doc.add_heading(document_title, level=0)
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta = doc.add_paragraph(f"Generated {generated}")
    meta.runs[0].italic = True

    for section in sections:
        for block in section.content_blocks or []:
            _add_block_to_docx(doc, block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
